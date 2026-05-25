from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate


OUT_DIR = Path("docs")
DOCX_PATH = OUT_DIR / "autora-koncepcijas-apraksts-labots.docx"
PDF_PATH = OUT_DIR / "autora-koncepcijas-apraksts-labots.pdf"
TXT_PATH = OUT_DIR / "autora-koncepcijas-apraksts-labots.txt"

TITLE = "Autora koncepcijas, vizuālā un tehniskā risinājuma apraksts"
WORK_TITLE = "Viena darba nedēļa datorikas studentes dzīvē"
AUTHOR = "Sabīne Stašūna"
SITE = "https://manalieliskanedela.vercel.app/#sakums"

SECTIONS = [
    (
        "Koncepcija un saturs",
        'Darba nosaukums ir "Viena darba nedēļa datorikas studentes dzīvē". Tīmekļa vietne publicēta adresē https://manalieliskanedela.vercel.app/#sakums, un tās mērķis ir parādīt vienu tipisku studiju nedēļu no datorikas studentes skatpunkta. Es izvēlējos vienas lapas struktūru, jo šāds formāts ļauj skatītājam secīgi iziet cauri visām nedēļas daļām: sākumam, lekciju grafikam, mājasdarbiem un projektiem, studiju platformām, ceļam uz universitāti un brīvajam laikam. Šīs sadaļas kopā veido pilnu priekšstatu par to, ka studijas nav tikai lekcijas, bet arī plānošana, pārvietošanās, digitālā komunikācija un atpūta.',
    ),
    (
        "Ilustratīvais materiāls",
        'Vizuālais risinājums veidots mierīgs un pārskatāms, lai saturs būtu viegli uztverams gan datorā, gan telefonā. Katrai sadaļai ir skaidrs nosaukums, paskaidrojošs teksts un ilustratīvs attēls. Lekciju grafikam izmantota tāfeles bilde ar uzrakstu "Test", jo tā asociējas ar pārbaudes darbiem, lekcijām un mācību spriedzi. Mājasdarbu sadaļā ievietots atvērtas grāmatas attēls, kas vizuāli papildina tēmu par pierakstiem un patstāvīgo darbu. Studiju platformu sadaļā izmantota tikšanās bilde ar datoriem, jo tā atgādina par tiešsaistes saziņu, Teams sarunām un kopīgu projektu darbu. Ceļa sadaļā ievietota bibliotēkas telpa, kas simboliski saistās ar universitātes vidi un nokļūšanu uz studijām. Brīvā laika sadaļā izmantota sporta zāles bilde, lai parādītu, ka pēc mācībām ir vajadzīga kustība un atslodze.',
    ),
    (
        "Grafiskie elementi un piekļūstamība",
        "Grafiskie elementi lapā ir izmantoti kā orientieri, nevis kā dekorācijas bez nozīmes. Ikonas pie sadaļu virsrakstiem palīdz ātri saprast, par ko būs konkrētā daļa, savukārt fotogrāfijas padara lapu dzīvāku un tuvāku reālai ikdienai. Krāsu izvēle ir atturīga: gaišajā režīmā dominē balts fons un tumšs teksts, bet tumšajā režīmā izmantots tumšs fons ar gaišu tekstu. Tas uzlabo lasāmību un atbilst vizuālās piekļūstamības prasībām. Attēliem pievienoti alternatīvie apraksti, lai saturs būtu saprotamāks arī tad, ja lietotājs izmanto ekrānlasītāju vai attēls netiek ielādēts.",
    ),
    (
        "Tehniskais risinājums",
        "Tehniski vietne veidota ar Next.js, React un Tailwind CSS. Šādu risinājumu izvēlējos tāpēc, ka tas ļauj sadalīt lapu komponentēs un uzturēt katru sadaļu atsevišķi. React komponentes padara struktūru pārskatāmu, bet Tailwind CSS palīdz ātri veidot responsīvu dizainu ar vienotu atstarpju, krāsu un izkārtojuma sistēmu. Lekciju grafiks ir izveidots kā ekspandējams saraksts, lai informācija neaizņemtu pārāk daudz vietas un lietotājs varētu atvērt tikai interesējošo dienu. Vietnē ir arī gaišais un tumšais režīms, ko iespējams pārslēgt navigācijas joslā.",
    ),
    (
        "Publicēšana un kopsavilkums",
        "Publicēšanai izmantota Vercel platforma, jo tā labi darbojas ar Next.js projektiem un ļauj lapu apskatīt pārlūkprogrammā bez reģistrēšanās vai papildu programmatūras instalēšanas. Projekta kods glabājas GitHub repozitorijā, tāpēc izmaiņas var ērti pārskatīt, labot un publicēt atkārtoti. Vietnē nav izmantoti LU vai EZTF logotipi, tādēļ netiek pārkāptas vizuālās identitātes vadlīnijas. Kopumā vietnes koncepcija balstās uz vienkāršu, personisku un saprotamu stāstu par studiju nedēļu, kur teksts, attēli un tehniskais risinājums kopā palīdz uztvert saturu bez liekas sarežģītības.",
    ),
]


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-zĀ-ž0-9]+(?:-[A-Za-zĀ-ž0-9]+)?", text))


def set_font(style, name: str, size: int, bold: bool = False) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_label_run(paragraph, label: str, value: str) -> None:
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(value)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    set_font(styles["Normal"], "Arial", 11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    set_font(styles["Title"], "Arial", 20)
    styles["Title"].paragraph_format.space_after = Pt(8)
    set_font(styles["Heading 1"], "Arial", 16, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run(TITLE)

    meta = doc.add_paragraph()
    add_label_run(meta, "Darbs: ", WORK_TITLE + "\n")
    add_label_run(meta, "Autore: ", AUTHOR + "\n")
    add_label_run(meta, "Publicētā vietne: ", SITE)

    for heading, body in SECTIONS:
        doc.add_heading(heading, level=1)
        paragraph = doc.add_paragraph(body)
        paragraph.paragraph_format.first_line_indent = Inches(0.25)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Sabīne Stašūna | koncepcijas apraksts")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.core_properties.author = AUTHOR
    doc.core_properties.title = TITLE
    doc.save(DOCX_PATH)


def esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf() -> None:
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    pdfmetrics.registerFont(TTFont("ArialLocal", str(font_path)))
    pdfmetrics.registerFont(TTFont("ArialLocal-Bold", str(bold_path)))

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LVTitle", fontName="ArialLocal-Bold", fontSize=16, leading=20, alignment=TA_LEFT, spaceAfter=8))
    styles.add(ParagraphStyle(name="LVMeta", fontName="ArialLocal", fontSize=10.5, leading=15, spaceAfter=12))
    styles.add(ParagraphStyle(name="LVHeading", fontName="ArialLocal-Bold", fontSize=12.5, leading=16, spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="LVBody", fontName="ArialLocal", fontSize=11, leading=16, firstLineIndent=0.6 * cm, spaceAfter=7))

    story = [
        Paragraph(esc(TITLE), styles["LVTitle"]),
        Paragraph(f"<b>Darbs:</b> {esc(WORK_TITLE)}<br/><b>Autore:</b> {esc(AUTHOR)}<br/><b>Publicētā vietne:</b> {esc(SITE)}", styles["LVMeta"]),
    ]
    for heading, body in SECTIONS:
        story.append(Paragraph(esc(heading), styles["LVHeading"]))
        story.append(Paragraph(esc(body), styles["LVBody"]))

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    body_text = "\n\n".join(body for _, body in SECTIONS)
    count = word_count(body_text)
    if not 400 <= count <= 700:
        raise ValueError(f"Word count must be 400-700, got {count}")
    TXT_PATH.write_text(
        f"{TITLE}\n\nDarbs: {WORK_TITLE}\nAutore: {AUTHOR}\nPublicētā vietne: {SITE}\n\n"
        + "\n\n".join(f"{heading}\n{body}" for heading, body in SECTIONS)
        + f"\n\nVārdu skaits: {count}\n",
        encoding="utf-8",
    )
    build_docx()
    build_pdf()
    print(f"Created {DOCX_PATH}")
    print(f"Created {PDF_PATH}")
    print(f"Word count: {count}")


if __name__ == "__main__":
    main()
